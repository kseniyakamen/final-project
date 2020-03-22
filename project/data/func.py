from docx import Document #library to read .doc and .docx files into the enviornment
import pandas as pd #data processing library
from nltk.stem.porter import PorterStemmer
from nltk.stem.wordnet import WordNetLemmatizer
import nltk
import sqlite3
#nltk.download('stopwords')
from nltk.corpus import stopwords
from nltk.stem.porter import PorterStemmer
from nltk.tokenize import RegexpTokenizer
#nltk.download('wordnet') 
from nltk.stem.wordnet import WordNetLemmatizer
from sklearn.feature_extraction.text import TfidfTransformer
import os
from os import path
from os import listdir
from os.path import isfile, join
from sklearn.feature_extraction.text import CountVectorizer
import re


def table_extractor(table):
    # creating dictionary that holds the values found in the table
    keys = None 
    for i, row in enumerate(table.rows):
        text = (cell.text for cell in row.cells)

        if i == 0:
            keys = tuple(text)
            continue
        row_data = dict(zip(keys, text))


    #manipulating the dictionary to allow easier and more comperhensive key word   
    skils_n_lang = [row[1] for row in row_data.items()]
    elements = []
    for row in skils_n_lang:
        row = row.split("\n")
        elements.append(row)
    final_skills = []
    for sublist in elements:
        for item in sublist:
            final_skills.append(item)
    return final_skills



def corpus_extractor(sub_set):
    
    corpus = []
    for i in sub_set['sentence']:
        #Remove punctuations
        text = re.sub('[^a-zA-Z]', ' ', i)

        text = re.sub(" \d+", " ", text)

        #Convert to lowercase
        text = text.lower()

        #remove tags
        text=re.sub("&lt;/?.*?&gt;"," &lt;&gt; ",text) #explain what re is - regular expressions

        # remove special characters and digits
        text=re.sub("(\\d|\\W)+"," ",text) #explain what re is - regular expressions

        ##Convert to list from string
        text = text.split()

        #Lemmatisation
        lem = WordNetLemmatizer() #explain what Lemmatization is. 
        text = " ".join(text) #join the different lists into a string
        corpus.append(text) #append the string to the corpus list

    if len(corpus) == 0:
        return []
    else:
        #creating a vector of word counts
        cv=CountVectorizer(max_df=0.8, max_features=10000, ngram_range=(1,3)) #CountVectorizer
        X=cv.fit_transform(corpus)
        
        #we transform the vector of word counts into a matrix of intergers
        tfidf_transformer=TfidfTransformer(smooth_idf=True,use_idf=True)
        tfidf_transformer.fit(X)

        # print idf values
        df_idf = pd.DataFrame(tfidf_transformer.idf_, index=cv.get_feature_names(),columns=["idf_weights"])

        # sort ascending
        return list(df_idf.index)
    
    
    
    
def converter(word_doc):
    
    #reading the .docx Document
    document = Document(word_doc)
    
    #extracting the paragraphs from the document object we just created
    test_list = [str(row.text.strip()) for row in document.paragraphs]

    #removing blank_spaces("")
    while("" in test_list): 
        test_list.remove("")

    name = test_list[0] #Grabbing the first line of the list which holds the name value
    
    #Converting the reminder of rows into a dataframe with two columns: 
        #one index[n = number of lines in the CV] 
        #one holding a sentence for that respective line
    dataset = pd.DataFrame(test_list[1:], columns= ["sentence"])
    

    #getting the index of the 3 different paragraph sections
    i = 0 #setting the iteration value to 0 
    dicts = dict() #creating an empty dictionary
    drop_idx = [] #creating an empty list
    
    for row in dataset["sentence"]: #starting a loop which looks through each row in the info column
        i += 1 #increase the iteration value by 1
        if row == "Previous Experience" or row == "Education and Certifications" or row == "Skills": #identifiying the rows at the start of each new section
            dicts[row] = i #getting the index for those precise rows
        elif row == "Key Achievements": #identifiying a common 
            drop_idx.append(i-1)
        else:
            pass
    
    idx = [i-1 for i in dicts.values()]
    dataset = dataset.drop(drop_idx)
    
    #checking if there are tables in the document object:
    lst_skills = []
    tables = document.tables
    if len(tables) >= 1: #if table is present we call table_extractor function
        for i in range(len(tables)):
            lst = table_extractor(tables[i]) #call table_extractor
            lst_skills.extend(lst)
    
    return dataset, idx , list(dicts.keys()), lst_skills, name



def corpus_redacter(dataset,idx, keys, table_skills,name):
    candidate = dict()
    word_lst = []
    
    #loop through the subsets
    for i in range(len(idx)):
        if i == 2:
            word_lst.extend(table_skills) #extend list of keywords with table skills if there are any
        else:
            sub_set = dataset[idx[i] :idx[i+1]] #extract sub-section thanks to idx list
            word_lst.extend(corpus_extractor(sub_set)) #run corpus_extractor function on subset onject

    word_lst = list(pd.Series(word_lst).drop_duplicates()) #removing duplicates
    candidate[name] = " ".join(word_lst) #join all the values of the list into a single string and append it to the candidate key = name
    return candidate


def sql_update_table_creator():
    
    path = os.getcwd() + "/cvs"
    onlyfiles = [f for f in listdir(path) if isfile(join(path, f))] #get only the text name of the files and their "." extentision

    #create three empty lists to hold the values we want to extract from each .docx document
    names = []
    emails = []
    keywords = []
    
    for row in onlyfiles:
        find = ".docx"
        if find in row:
            dataset, idx, keys, table_skills, name =  converter("cvs/"+row)
            candidate = corpus_redacter(dataset, idx, keys, table_skills,name)
            person = list(candidate.keys())[0]
            names.append(person)
            email = person.replace(" ", ".").lower() + "@replgroup.com"
            emails.append(email)
            keyword = candidate[name]
            keywords.append(keyword)
        else:
            pass
        
    df = pd.DataFrame(names, columns = ["name"])
    df["email"] = emails
    df["keywords"] = keywords
    
    return df
  

